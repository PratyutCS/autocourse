import json
import sys
import os
import uuid
from filelock import FileLock, Timeout
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx2pdf import convert
from PyPDF2 import PdfMerger

# Read data from command line (expected to be a JSON string)
input_data = sys.stdin.buffer.read().decode('utf-8')
data = json.loads(input_data)

if data.get('Program'):
    program_value = data.get("Program", "0")
    proram_number = str(program_value)
    if program_value == "0" or program_value > 10:
        with open('./extractor/sample.docx', 'rb') as f:
            doc = Document(f)
    else:
        with open('./extractor/'+proram_number+'.docx', 'rb') as f:
            doc = Document(f)
else:
    with open('./extractor/sample.docx', 'rb') as f:
        doc = Document(f)

header = ""

def rep(doc, key):
    placeholder = f'{{{{{key}}}}}'
    value = data.get(key, "")
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            text = paragraph.text.replace(placeholder, value)
            paragraph.clear()
            if ':' in text:
                before_colon, after_colon = text.split(':', 1)
                run1 = paragraph.add_run(before_colon + ':')
                run1.bold = True
                run1.font.size = Pt(14)
                run2 = paragraph.add_run(after_colon)
                run2.font.size = Pt(14)
            else:
                run = paragraph.add_run(text)
                run.font.size = Pt(14)

header = ""

if data.get('course_name'):
    rep(doc, "course_name")
    header += data.get("course_name", "") + " ("
if data.get('course_code'):
    rep(doc, "course_code")
    header += data.get("course_code", "") + ") Sem: "
if data.get('Module/Semester'):
    rep(doc, "Module/Semester")
    header += data.get("Module/Semester", "")
if data.get('Session'):
    rep(doc, "Session")

#######################################################################################################################
# Code for adding headers to the document

section = doc.sections[0]
section.different_first_page_header_footer = True

# First page header
first_page_header = section.first_page_header
first_page_paragraph = first_page_header.paragraphs[0] if first_page_header.paragraphs else first_page_header.add_paragraph()
first_page_paragraph.text = header
first_page_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
first_page_paragraph.paragraph_format.right_indent = Inches(0.2)

# Header for the rest of the pages
default_header = section.header
default_paragraph = default_header.paragraphs[0] if default_header.paragraphs else default_header.add_paragraph()
default_paragraph.text = header
default_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
default_paragraph.paragraph_format.right_indent = Inches(0.2)




#######################################################################################################################
# Course Description and Objectives
if data.get('course_description'):
    doc.add_page_break()
    course_heading = doc.add_heading(level=1)
    course_heading.paragraph_format.left_indent = Inches(0.5)
    course_run = course_heading.add_run('6. Course Description and its objectives')
    course_run.font.name = 'Carlito'
    course_run.font.size = Pt(16)
    course_run.font.color.rgb = RGBColor(28, 132, 196)
    course_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # empty line after heading

    course_description = data['course_description'].encode('utf-8', 'ignore').decode('utf-8')

    for line in course_description.split('\n'):
        if line.strip():  # skip empty lines
            paragraph = doc.add_paragraph(line.strip())
            paragraph.paragraph_format.left_indent = Inches(0.7)
            paragraph.paragraph_format.right_indent = Inches(0.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


#######################################################################################################################
# CO-PO Mapping Section
if data.get('copoMappingData'):
    doc.add_page_break()
    co_heading = doc.add_heading(level=1)
    co_run = co_heading.add_run('7. Course Outcomes and CO-PO Mapping')
    co_run.font.name = 'Carlito'
    co_run.font.size = Pt(16)
    co_run.font.color.rgb = RGBColor(28, 132, 196)
    co_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if data['copoMappingData'].get('courseOutcomes'):
        doc.add_paragraph()
        course_outcomes_heading = doc.add_paragraph()
        course_outcomes_run = course_outcomes_heading.add_run('Course Outcomes:')
        course_outcomes_run.bold = True
        course_outcomes_run.font.size = Pt(12)
        course_outcomes_heading.paragraph_format.left_indent = Inches(0.7)
        course_outcomes_heading.paragraph_format.right_indent = Inches(0.5)
        course_outcomes_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for co_key, co_value in data['copoMappingData']['courseOutcomes'].items():
            if isinstance(co_value, dict):
                doc.add_paragraph()
                paragraph = doc.add_paragraph(f'{co_key}: {co_value["description"]}')
                paragraph.paragraph_format.left_indent = Inches(0.7)
                paragraph.paragraph_format.right_indent = Inches(0.5)
                for bullet in co_value.get('bullets', []):
                    bullet_paragraph = doc.add_paragraph('• ' + bullet)
                    bullet_paragraph.paragraph_format.left_indent = Inches(1)
                    bullet_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                doc.add_paragraph()
                paragraph = doc.add_paragraph(f'{co_key}: {co_value}')
                paragraph.paragraph_format.left_indent = Inches(0.7)
                paragraph.paragraph_format.right_indent = Inches(0.5)

    if data['copoMappingData'].get('mappingData'):
        doc.add_paragraph()
        mapping_heading = doc.add_paragraph()
        mapping_run = mapping_heading.add_run('CO/PO Mapping:')
        mapping_run.bold = True
        mapping_run.font.size = Pt(10)
        mapping_heading.paragraph_format.left_indent = Inches(0.7)
        mapping_heading.paragraph_format.right_indent = Inches(0.5)
        mapping_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        co_po_mapping = data['copoMappingData']['mappingData']
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=len(co_po_mapping[list(co_po_mapping.keys())[0]]) + 1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Course Outcomes (CO)'
        for i, po_key in enumerate(co_po_mapping[list(co_po_mapping.keys())[0]].keys()):
            hdr_cells[i + 1].text = po_key

        for co_key, po_values in co_po_mapping.items():
            row_cells = table.add_row().cells
            row_cells[0].text = co_key
            for i, value in enumerate(po_values.values()):
                row_cells[i + 1].text = str(value) if value != "" else ""

        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(0.06)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        # Set borders for the table cells
        def set_cell_border(cell, border_type, border_size, border_color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:color'), border_color)
            element = tcPr.xpath(f"./w:{border_type}")
            if element:
                element[0].getparent().replace(element[0], border)
            else:
                tcPr.append(border)

        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, 'top', 4, '000000')
                set_cell_border(cell, 'bottom', 4, '000000')
                set_cell_border(cell, 'left', 4, '000000')
                set_cell_border(cell, 'right', 4, '000000')

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)

#######################################################################################################################
# Course Syllabus Section
if data.get('Course Syllabus'):
    doc.add_page_break()
    syllabus_heading = doc.add_heading(level=1)
    syllabus_run = syllabus_heading.add_run('8. Detailed Session wise Plan & Course Syllabus')
    syllabus_run.font.name = 'Carlito'
    syllabus_run.font.size = Pt(16)
    syllabus_run.font.color.rgb = RGBColor(28, 132, 196)
    syllabus_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    headers = ['Sr. No.', 'Content', 'CO', 'Sessions']
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)

    for item in data['Course Syllabus']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get('srNo', ''))
        row_cells[1].text = str(item.get('content', ''))
        row_cells[2].text = str(item.get('co', ''))
        row_cells[3].text = str(item.get('sessions', ''))
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    widths = [Inches(0.8), Inches(3.5), Inches(0.8), Inches(0.8)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = widths[idx]

    def set_cell_border(cell, border_type, border_size, border_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        element = tcPr.xpath(f"./w:{border_type}")
        if element:
            element[0].getparent().replace(element[0], border)
        else:
            tcPr.append(border)

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 4, '000000')
            set_cell_border(cell, 'bottom', 4, '000000')
            set_cell_border(cell, 'left', 4, '000000')
            set_cell_border(cell, 'right', 4, '000000')

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
    

doc.add_paragraph()
doc.add_paragraph()

#######################################################################################################################
# Learning Resources Section
def create_learning_resources_doc(data):
    # doc.add_page_break()
    timetable_heading = doc.add_heading(level=1)
    timetable_run = timetable_heading.add_run('Learning Resources')
    timetable_run.font.name = 'Carlito'
    timetable_run.font.size = Pt(12)
    # timetable_run.font.color.rgb = RGBColor(28, 132, 196)
    timetable_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    textbooks_heading = doc.add_heading('Text Books:', level=2)
    textbooks_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    textbooks_run = textbooks_heading.runs[0]
    textbooks_run.bold = True
    textbooks_run.font.size = Pt(12)
    
    if data.get('textBooks'):
        for book in data['textBooks']:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(1)
            bullet = para.add_run("✓ ")
            bullet.font.size = Pt(12)
            # Clean the book text to remove any problematic unicode characters.
            safe_book = book.encode('utf-8', 'replace').decode('utf-8')
            text = para.add_run(safe_book)
            text.font.size = Pt(12)

    doc.add_paragraph()
    
    ref_heading = doc.add_heading('Reference Links:', level=2)
    ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ref_run = ref_heading.runs[0]
    ref_run.bold = True
    ref_run.font.size = Pt(12)
    
    if data.get('referenceLinks'):
        for link in data['referenceLinks']:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(1)
            bullet = para.add_run("• ")
            bullet.font.size = Pt(12)
            # Clean the link text similarly.
            safe_link = link.encode('utf-8', 'replace').decode('utf-8')
            link_run = para.add_run(safe_link)
            link_run.font.size = Pt(12)
            link_run.font.color.rgb = RGBColor(0, 0, 255)
            link_run.underline = True

create_learning_resources_doc(data["Learning Resources"])



#######################################################################################################################
# Weekly Timetable Section
if data.get('weeklyTimetableData'):
    doc.add_page_break()
    timetable_heading = doc.add_heading(level=1)
    timetable_run = timetable_heading.add_run('9. Weekly Timetable')
    timetable_run.font.name = 'Carlito'
    timetable_run.font.size = Pt(16)
    timetable_run.font.color.rgb = RGBColor(28, 132, 196)
    timetable_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    
    # Check if this is an MBA timetable (has 'entries' field) or regular timetable
    if 'entries' in data['weeklyTimetableData']:
        # MBA Timetable - Format entries in a different layout
        entries = data['weeklyTimetableData'].get('entries', [])
        
        if entries:
            # Sort entries by day and time for better organization
            days_order = {
                'Monday': 1, 'Tuesday': 2, 'Wednesday': 3, 'Thursday': 4, 
                'Friday': 5, 'Saturday': 6, 'Sunday': 7
            }
            
            # Sort entries by day first, then by time (AM before PM, then by hour)
            sorted_entries = sorted(entries, key=lambda x: (
                days_order.get(x.get('day', ''), 999),
                0 if x.get('period', '') == 'AM' else 1,
                x.get('hour', 0)
            ))
            
            # Create a table for the MBA timetable with columns: Day, Start Time, End Time, Duration
            table = doc.add_table(rows=len(sorted_entries) + 1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set header row
            header_cells = table.rows[0].cells
            headers = ['Day', 'Start Time', 'End Time', 'Duration (hrs)']
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                paragraph = header_cells[i].paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(header)
                run.bold = True
                run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Helper function to calculate end time
            def get_end_time(start_hour, period, duration):
                hour = start_hour
                end_period = period
                
                hour_part = int(duration)
                minute_part = int((duration - hour_part) * 60)
                
                hour += hour_part
                minutes = minute_part
                
                # Handle period change
                if hour >= 12:
                    if hour > 12:
                        hour = hour - 12
                    if period == 'AM':
                        end_period = 'PM'
                
                # Format the time
                return f"{hour}:{minutes:02d} {end_period}"
            
            # Fill the table with MBA timetable data
            for i, entry in enumerate(sorted_entries):
                row_cells = table.rows[i + 1].cells
                day = entry.get('day', '')
                hour = entry.get('hour', 0)
                period = entry.get('period', 'AM')
                duration = entry.get('duration', 1)
                
                # Format the start and end times
                start_time = f"{hour}:00 {period}"
                end_time = get_end_time(hour, period, duration)
                
                # Set cell values
                row_cells[0].text = day
                row_cells[1].text = start_time
                row_cells[2].text = end_time
                row_cells[3].text = str(duration)
                
                # Apply formatting to cells
                for j, cell in enumerate(row_cells):
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            
            # Set column widths for better readability
            column_widths = [Inches(1.5), Inches(1.2), Inches(1.2), Inches(1)]
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx < len(column_widths):
                        cell.width = column_widths[idx]
            
            # Add a visual weekly overview table below the entries table
            doc.add_paragraph().add_run().add_break()
            overview_para = doc.add_paragraph()
            overview_run = overview_para.add_run("Weekly Overview")
            overview_run.bold = True
            overview_run.font.size = Pt(12)
            
            # Create weekly overview table
            days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            overview_table = doc.add_table(rows=1, cols=len(days))
            overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add days as headers
            header_cells = overview_table.rows[0].cells
            for i, day in enumerate(days):
                header_cells[i].text = day
                paragraph = header_cells[i].paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(day)
                run.bold = True
                run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add a row for classes
            class_row = overview_table.add_row()
            class_cells = class_row.cells
            
            # Group entries by day
            day_entries = {}
            for day in days:
                day_entries[day] = [e for e in sorted_entries if e.get('day') == day]
            
            # Add class entries to each day cell
            for i, day in enumerate(days):
                cell = class_cells[i]
                entries_for_day = day_entries[day]
                
                if entries_for_day:
                    for entry in entries_for_day:
                        start_time = f"{entry.get('hour', 0)}:00 {entry.get('period', 'AM')}"
                        end_time = get_end_time(entry.get('hour', 0), entry.get('period', 'AM'), entry.get('duration', 1))
                        
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run(f"{start_time} - {end_time}")
                        run.font.size = Pt(9)
                        run.bold = True
                        
                        detail_para = cell.add_paragraph()
                        detail_run = detail_para.add_run(f"{data['course_code']} ({entry.get('duration')}hr)")
                        detail_run.font.size = Pt(8)
                        detail_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add spacing between entries
                        if entry != entries_for_day[-1]:
                            cell.add_paragraph()
                else:
                    paragraph = cell.add_paragraph()
                    run = paragraph.add_run("No class")
                    run.font.size = Pt(9)
                    run.italic = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set even height and width for overview table
            column_width = Inches(1.1)
            for row in overview_table.rows:
                for cell in row.cells:
                    cell.width = column_width
                    
    else:
        # Regular Timetable (Original code for non-MBA timetable)
        time_slots = list(data['weeklyTimetableData']['Monday'].keys())
        days = list(data['weeklyTimetableData'].keys())
        table = doc.add_table(rows=len(time_slots) + 1, cols=len(days) + 1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Time'
        for i, day in enumerate(days):
            header_cells[i + 1].text = day
            paragraph = header_cells[i + 1].paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(day)
            run.bold = True
            run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        time_paragraph = header_cells[0].paragraphs[0]
        time_run = time_paragraph.runs[0] if time_paragraph.runs else time_paragraph.add_run('Time')
        time_run.bold = True
        time_run.font.size = Pt(10)
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, time_slot in enumerate(time_slots):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = time_slot
            for j, day in enumerate(days):
                is_available = data['weeklyTimetableData'][day][time_slot]
                cell = row_cells[j + 1]
                cell.text = data['course_name'] + " (" + data['course_code'] + ")" if is_available else ''
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        column_widths = [Inches(1.1)] + [Inches(1)] * len(days)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(column_widths):
                    cell.width = column_widths[idx]

    # Apply border styling to the table (works for both table types)
    def set_cell_border(cell, border_type, border_size, border_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        element = tcPr.xpath(f"./w:{border_type}")
        if element:
            element[0].getparent().replace(element[0], border)
        else:
            tcPr.append(border)

    # Apply borders to all tables in the timetable section
    for table in doc.tables[-2:]:  # Apply to the last two tables (entries and overview for MBA, or just one table for regular)
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, 'top', 4, '000000')
                set_cell_border(cell, 'bottom', 4, '000000')
                set_cell_border(cell, 'left', 4, '000000')
                set_cell_border(cell, 'right', 4, '000000')

            # Prevent row splitting across pages
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)  


################################################## StudentData ########################################################

if data.get('studentData'):
    doc.add_page_break()
    reg_students_heading = doc.add_heading(level=1)
    reg_students_run = reg_students_heading.add_run('10. Registered Students List')
    reg_students_run.font.name = 'Carlito'
    reg_students_run.font.size = Pt(16)
    reg_students_run.font.color.rgb = RGBColor(28, 132, 196)
    reg_students_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    
    # Get student data and assessment headers
    student_data = data['studentData']['data']
    # assessment_keys = list(data['studentData']['maxMarks'].keys())
    
    # Get actual header names from the first student entry
    all_keys = list(student_data[0].keys())
    
    # First 3 keys are student info, remainder are assessment data
    student_info_keys = all_keys[:3]  # Position-based access instead of hardcoded names
    
    # Combine fixed headers with assessment headers
    all_headers = student_info_keys 
    
    # Create the table
    table = doc.add_table(rows=1, cols=len(all_headers))
    
    # Set header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(all_headers):
        header_cells[i].text = header_text
        paragraph = header_cells[i].paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(header_text)
        run.bold = True
        run.font.size = Pt(9)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add student rows
    for student in student_data:
        row_cells = table.add_row().cells
        
        # Add student info columns by position from keys
        for i, key in enumerate(student_info_keys):
            row_cells[i].text = str(student.get(key, ''))
        
        # Add assessment values
        
        
        # Format cell text
        for cell in row_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
    
    # Calculate dynamic column widths
    total_width = Inches(6.3)  # Total table width
    column_width = total_width / len(all_headers)
    
    # Apply formatting
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = column_width
    
    # Set table borders
    def set_cell_border(cell, border_type, border_size, border_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        element = tcPr.xpath(f"./w:{border_type}")
        if element:
            element[0].getparent().replace(element[0], border)
        else:
            tcPr.append(border)
    
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 3, '000000')
            set_cell_border(cell, 'bottom', 3, '000000')
            set_cell_border(cell, 'left', 3, '000000')
            set_cell_border(cell, 'right', 3, '000000')
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ensure rows don't split across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)


###############################################  internalAssessmentData  ##############################################################

if data.get('internalAssessmentData') and data['internalAssessmentData'].get('components'):
    # Add a page break and heading for Internal Assessment Data
    doc.add_page_break()
    assessment_heading = doc.add_heading(level=1)
    assessment_run = assessment_heading.add_run('11. Internal Assessment Data')
    assessment_run.font.name = 'Carlito'
    assessment_run.font.size = Pt(16)
    assessment_run.font.color.rgb = RGBColor(28, 132, 196)
    assessment_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # Extract the components dictionary
    components = data['internalAssessmentData']['components']

    # Generate headers dynamically from the first component
    headers = list(next(iter(components.values())).keys())

    # Create a table with the number of columns based on headers
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Populate header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header.capitalize()  # Capitalize for better formatting
        # Style header cells
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Populate rows with component data
    for component_id, component_data in components.items():
        row_cells = table.add_row().cells
        for i, key in enumerate(headers):
            row_cells[i].text = str(component_data.get(key, ''))
            # Style cells
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Set column widths (adjustable based on content)
    # Increase width of the last column
    column_widths = [Inches(1)] * (len(headers) - 1) + [Inches(2.5)]  # Last column gets 2 inches
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(column_widths):
                cell.width = column_widths[idx]

    # Apply borders to cells
    def set_cell_border(cell, border_type, border_size, border_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)

        element = tcPr.xpath(f"./w:{border_type}")
        if element:
            element[0].getparent().replace(element[0], border)
        else:
            tcPr.append(border)

    # Apply borders to all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 4, '000000')
            set_cell_border(cell, 'bottom', 4, '000000')
            set_cell_border(cell, 'left', 4, '000000')
            set_cell_border(cell, 'right', 4, '000000')

    # Prevent table rows from splitting across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

################################### to add 12 ###########################################################

###################################### create_learner_categorization_partial_sem ###################################################################################

def format_cell(cell, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Format a table cell"""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    
    for run in paragraph.runs:
        run.font.bold = bold
        run.font.size = Pt(10)
    
    # If no runs, create one
    if not paragraph.runs:
        run = paragraph.add_run(cell.text)
        run.font.bold = bold
        run.font.size = Pt(10)
        cell.text = ""  # Clear text since we've added it to the run

def set_cell_border(cell, border_type, border_size, border_color):
    """Set cell border properties"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    border = OxmlElement(f'w:{border_type}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(border_size))
    border.set(qn('w:color'), border_color)
    element = tcPr.xpath(f"./w:{border_type}")
    if element:
        element[0].getparent().replace(element[0], border)
    else:
        tcPr.append(border)

def set_table_column_widths(table, widths):
    """Set column widths for a table"""
    for i, width in enumerate(widths):
        for row in table.rows:
            if i < len(row.cells):
                row.cells[i].width = width

def prevent_table_row_breaks(table):
    """Ensure table rows don't split across pages"""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
def create_learner_categorization_partial_sem(doc, data):
    # Add page break and section heading
    doc.add_page_break()
    heading = doc.add_heading(level=1)
    run = heading.add_run('13. Sample Evaluated Internal Submissions and Identification of weak students.')
    run.font.name = 'Carlito'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(28, 132, 196)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    
    # Get full student data and learner categories from input data
    student_data = data.get("studentData", {}).get("data", [])
    # "par_sem_slowLearner" is a list of two lists: first for advanced, second for slow learners
    par_sem = data.get("par_sem_slowLearner", [])
    advanced_ids = {stud["id"] for stud in par_sem[0]} if len(par_sem) > 0 else set()
    slow_ids = {stud["id"] for stud in par_sem[1]} if len(par_sem) > 1 else set()
    
    # Classify each student based on their Unique Id.
    advanced_learners = []
    slow_learners = []
    medium_learners = []
    for student in student_data:
        unique_id = student.get("Unique Id.")
        if unique_id in advanced_ids:
            advanced_learners.append(student)
        elif unique_id in slow_ids:
            slow_learners.append(student)
        else:
            medium_learners.append(student)
            
    # 1. Create summary table with updated heading
    summary_heading = doc.add_heading('Learner Categories Summary for Partial Semester', level=2)
    summary_heading.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.autofit = False
    
    # Header row
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Learner Category"
    header_cells[1].text = "Number of Students"
    format_cell(header_cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(header_cells[1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    categories = [
        ("Advanced Learners", len(advanced_learners)),
        ("Medium Learners", len(medium_learners)),
        ("Low Performers", len(slow_learners))
    ]
    for idx, (category, count) in enumerate(categories):
        row = summary_table.rows[idx + 1]
        row.cells[0].text = category
        row.cells[1].text = str(count)
        format_cell(row.cells[0], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(row.cells[1], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
    col_widths = [Inches(3), Inches(3)]
    set_table_column_widths(summary_table, col_widths)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.autofit = False
    for row in summary_table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 4, '000000')
            set_cell_border(cell, 'bottom', 4, '000000')
            set_cell_border(cell, 'left', 4, '000000')
            set_cell_border(cell, 'right', 4, '000000')
            
    for row in summary_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    prevent_table_row_breaks(summary_table)
    doc.add_paragraph().add_run().add_break()
    
    # 2. Create detailed student table with appropriate color coding and updated heading
    detailed_heading = doc.add_heading('Student Learning Classification for Partial Semester', level=2)
    detailed_heading.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    
    total_students = len(student_data)
    student_table = doc.add_table(rows=total_students + 1, cols=2)
    student_table.autofit = False
    
    # Table header
    header_cells = student_table.rows[0].cells
    header_cells[0].text = "Student Name"
    header_cells[1].text = "Category"
    format_cell(header_cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(header_cells[1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_column_widths(student_table, [Inches(3), Inches(3)])
    
    # Combine students: advanced, medium, then slow
    all_students = advanced_learners + medium_learners + slow_learners
    for idx, student in enumerate(all_students):
        row = student_table.rows[idx + 1]
        name = student.get("Student Name", "")
        unique_id = student.get("Unique Id.")
        if unique_id in advanced_ids:
            category = "Advanced Learner"
            color = "C6E0B4"  # green
        elif unique_id in slow_ids:
            category = "Slow Learner"
            color = "FFEB9C"  # yellow
        else:
            category = "Medium Learner"
            color = "F2F2F2"  # grey

        row.cells[0].text = name
        row.cells[1].text = category
        format_cell(row.cells[0], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(row.cells[1], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), color)
            existing_shading = tc_pr.find(qn('w:shd'))
            if existing_shading is not None:
                tc_pr.remove(existing_shading)
            tc_pr.append(shading)
    
    student_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    student_table.autofit = False
    for row in student_table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 4, '000000')
            set_cell_border(cell, 'bottom', 4, '000000')
            set_cell_border(cell, 'left', 4, '000000')
            set_cell_border(cell, 'right', 4, '000000')
    for row in student_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    prevent_table_row_breaks(student_table)
    doc.add_paragraph()
    
# Call this function after create_co_attainment_analysis(doc, data) and before create_actions_doc(data)
create_learner_categorization_partial_sem(doc, data)


###################################################################################################################

def create_reflection_doc(data):
    if data.get('reflectionData'):
        # Add the main heading with a page break.
        doc.add_page_break()
        timetable_heading = doc.add_heading(level=1)
        timetable_run = timetable_heading.add_run(
            '14. Reflections from the Mid-term semester feedback received, and interventions made to enhance the student learning and continuous improvement in teaching and learning strategies.'
        )
        timetable_run.font.name = 'Carlito'
        timetable_run.font.size = Pt(16)
        timetable_run.font.color.rgb = RGBColor(28, 132, 196)
        timetable_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()

        for action in data['reflectionData']:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Inches(0.6)
            para.paragraph_format.first_line_indent = Pt(0)
            bullet_run = para.add_run("• ")
            bullet_run.font.size = Pt(12)
            action_run = para.add_run(action)
            action_run.font.size = Pt(12)

create_reflection_doc(data)


#########################################################################################################################
# Updated Actions Taken for Weak Students Section (create_actions_doc)
def create_actions_doc(data):
    if data.get('actionsForWeakStudentsData'):
        # Add the main heading with a page break.
        doc.add_page_break()
        timetable_heading = doc.add_heading(level=1)
        timetable_run = timetable_heading.add_run('15. Actions taken for low performers')
        timetable_run.font.name = 'Carlito'
        timetable_run.font.size = Pt(16)
        timetable_run.font.color.rgb = RGBColor(28, 132, 196)
        timetable_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()
        
        # Loop through each action and add as a bullet.
        for action in data['actionsForWeakStudentsData']:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Inches(0.6)
            para.paragraph_format.first_line_indent =  Pt(0)
            bullet_run = para.add_run("• ")
            bullet_run.font.size = Pt(12)
            action_run = para.add_run(action)
            action_run.font.size = Pt(12)

create_actions_doc(data)   

####################################### to add 16 ############################################################################

###################################################################################################################
        

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
if data.get('studentData'):

    # Get student data and find the attendance column
    student_data = data['studentData']['data']
    first_student = student_data[0]


     #2. DETAIL OF MARKS TABLE
######################################## 17 ################################################
    # doc.add_page_break()
    marks_heading = doc.add_heading(level=1)
    marks_run = marks_heading.add_run('17. Details of Marks in all components up to the End Semester including the grades')
    marks_run.font.name = 'Carlito'
    marks_run.font.size = Pt(16)
    marks_run.font.color.rgb = RGBColor(28, 132, 196)
    marks_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    
    # First 3 columns are always student info
    student_info_keys = list(first_student.keys())[:3]
    
    # Get assessment columns from maxMarks, exclude "TOTAL Marks"
    assessment_keys = [key for key in data['studentData']['maxMarks'].keys() 
                       if not key.lower().startswith('total')]
    
    # Find grading column if it exists
    grading_column = None
    for key in all_keys:
        if 'grade' in key.lower() or 'grading' in key.lower():
            grading_column = key
            break
    
    # Headers for the marks table: student info + assessments + (grade if exists)
    marks_headers = student_info_keys + assessment_keys
    if grading_column and grading_column not in marks_headers:
        marks_headers.append(grading_column)
    
    # Create table
    table = doc.add_table(rows=1, cols=len(marks_headers))
    
    # Set header row
    header_cells = table.rows[0].cells
    
    for i, header_text in enumerate(marks_headers):
        header_cells[i].text = header_text
        
        # Add max marks info for assessment columns
        if header_text in data['studentData']['maxMarks']:
            max_mark = data['studentData']['maxMarks'][header_text]
            header_cells[i].add_paragraph(f"Out of ({max_mark})")
            
        # Format all header text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add student marks rows
    for student in student_data:
        row_cells = table.add_row().cells
        
        # Add data for each header
        for i, header in enumerate(marks_headers):
            value = student.get(header, '')
            row_cells[i].text = str(value)
        
        # Format cell text
        for cell in row_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
    
    # Calculate dynamic column widths
    total_width = Inches(7.0)
    
    # Student info columns get 40% of width
    id_cols_count = len(student_info_keys)
    id_cols_total = 0.4 * total_width  
    id_col_width = id_cols_total / id_cols_count if id_cols_count > 0 else 0
    id_widths = [id_col_width] * id_cols_count
    
    # Assessment columns get 50% of width
    assessment_count = len(assessment_keys)
    assessment_total = 0.5 * total_width
    assessment_width = assessment_total / assessment_count if assessment_count > 0 else 0
    assessment_widths = [assessment_width] * assessment_count
    
    # Grade column gets 10% if it exists
    grade_width = []
    if grading_column in marks_headers:
        grade_width = [0.1 * total_width]
    
    widths = id_widths + assessment_widths + grade_width
    
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = widths[idx]
    
    # Set table borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 3, '000000')
            set_cell_border(cell, 'bottom', 3, '000000')
            set_cell_border(cell, 'left', 3, '000000')
            set_cell_border(cell, 'right', 3, '000000')
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ensure table rows don't split across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
    
    ########################################################################################
    def format_cell(cell, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """Format a table cell"""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        
        for run in paragraph.runs:
            run.font.bold = bold
            run.font.size = Pt(10)
        
        # If no runs, create one
        if not paragraph.runs:
            run = paragraph.add_run(cell.text)
            run.font.bold = bold
            run.font.size = Pt(10)
            cell.text = ""  # Clear text since we've added it to the run

    def set_cell_border(cell, border_type, border_size, border_color):
        """Set cell border properties"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        element = tcPr.xpath(f"./w:{border_type}")
        if element:
            element[0].getparent().replace(element[0], border)
        else:
            tcPr.append(border)

    def set_table_column_widths(table, widths):
        """Set column widths for a table"""
        for i, width in enumerate(widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = width

    def prevent_table_row_breaks(table):
        """Ensure table rows don't split across pages"""
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)

    #########################################################################################################################
    def create_learner_categorization(doc, data):
        # Add page break and section heading
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('18. Identification of advanced learners and low performers conducted at the end of the semester')
        run.font.name = 'Carlito'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(28, 132, 196)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        
        # Extract full student data and learner categories from input data
        student_data = data.get("studentData", {}).get("data", [])
        learner_categories = data.get("learnerCategories", [])
        
        # Identify student IDs in the learnerCategories arrays:
        advanced_ids = {student["id"] for student in learner_categories[0]} if len(learner_categories) > 0 else set()
        slow_ids = {student["id"] for student in learner_categories[1]} if len(learner_categories) > 1 else set()
        
        # Categorize each student from student_data into advanced, slow or medium
        advanced_learners = []
        slow_learners = []
        medium_learners = []
        for student in student_data:
            unique_id = student.get("Unique Id.")
            if unique_id in advanced_ids:
                advanced_learners.append(student)
            elif unique_id in slow_ids:
                slow_learners.append(student)
            else:
                medium_learners.append(student)
                
        # 1. Create summary table
        summary_heading = doc.add_heading('Learner Categories Summary', level=2)
        summary_heading.runs[0].font.size = Pt(14)
        doc.add_paragraph()
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.autofit = False
        
        # Header row
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = "Learner Category"
        header_cells[1].text = "Number of Students"
        format_cell(header_cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(header_cells[1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        categories = [
            ("Advanced Learners", len(advanced_learners)),
            ("Medium Learners", len(medium_learners)),
            ("Slow Learners", len(slow_learners))
        ]
        for idx, (category, count) in enumerate(categories):
            row = summary_table.rows[idx + 1]
            row.cells[0].text = category
            row.cells[1].text = str(count)
            format_cell(row.cells[0], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            format_cell(row.cells[1], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
        col_widths = [Inches(3), Inches(3)]
        set_table_column_widths(summary_table, col_widths)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        summary_table.autofit = False
        for row in summary_table.rows:
            for cell in row.cells:
                set_cell_border(cell, 'top', 4, '000000')
                set_cell_border(cell, 'bottom', 4, '000000')
                set_cell_border(cell, 'left', 4, '000000')
                set_cell_border(cell, 'right', 4, '000000')
        # Reduce font size
        for row in summary_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        prevent_table_row_breaks(summary_table)
        doc.add_paragraph().add_run().add_break()
        
        # 2. Create detailed student table with appropriate color coding
        detailed_heading = doc.add_heading('Student Learning Classification', level=2)
        detailed_heading.runs[0].font.size = Pt(14)
        doc.add_paragraph()
        
        total_students = len(student_data)
        student_table = doc.add_table(rows=total_students + 1, cols=2)
        student_table.autofit = False
        
        # Header
        header_cells = student_table.rows[0].cells
        header_cells[0].text = "Student Name"
        header_cells[1].text = "Category"
        format_cell(header_cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(header_cells[1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        col_widths = [Inches(3), Inches(3)]
        set_table_column_widths(student_table, col_widths)
        
        # Combine students in order: advanced, medium, then slow
        all_students = advanced_learners + medium_learners + slow_learners
        for idx, student in enumerate(all_students):
            row = student_table.rows[idx + 1]
            name = student.get("Student Name", "")
            unique_id = student.get("Unique Id.")
            if unique_id in advanced_ids:
                category = "Advanced Learner"
                color = "C6E0B4"  # green
            elif unique_id in slow_ids:
                category = "Slow Learner"
                color = "FFEB9C"  # yellow
            else:
                category = "Medium Learner"
                color = "F2F2F2"  # grey

            row.cells[0].text = name
            row.cells[1].text = category
            format_cell(row.cells[0], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            format_cell(row.cells[1], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Apply shading (color fill) to each cell in the row
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), color)
                existing_shading = tc_pr.find(qn('w:shd'))
                if existing_shading is not None:
                    tc_pr.remove(existing_shading)
                tc_pr.append(shading)
        
        student_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        student_table.autofit = False
        for row in student_table.rows:
            for cell in row.cells:
                set_cell_border(cell, 'top', 4, '000000')
                set_cell_border(cell, 'bottom', 4, '000000')
                set_cell_border(cell, 'left', 4, '000000')
                set_cell_border(cell, 'right', 4, '000000')
        for row in student_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        prevent_table_row_breaks(student_table)
        doc.add_paragraph()
        
    # Add this function call after create_co_attainment_analysis(doc, data) and before create_actions_doc(data)
    create_learner_categorization(doc, data)
    # 1. ATTENDANCE REPORT TABLE  
######################################## attendance report ###############################################
    doc.add_page_break()
    attendance_heading = doc.add_heading(level=1)
    attendance_run = attendance_heading.add_run('19. Attendance Report')
    attendance_run.font.name = 'Carlito'
    attendance_run.font.size = Pt(16)
    attendance_run.font.color.rgb = RGBColor(28, 132, 196)
    attendance_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    
    # Get all student keys
    all_keys = list(first_student.keys())
    
    # First 3 keys are always student info (Sr.No, ID, Name)
    student_info_keys = all_keys[:3]
    
    # Find attendance column by name (should be the last non-assessment column)
    attendance_column = None
    for key in all_keys:
        if 'attendance' in key.lower():
            attendance_column = key
            break
    
    # If no explicit attendance column found, use the last non-assessment column
    if not attendance_column:
        assessment_keys = list(data['studentData']['maxMarks'].keys())
        for key in reversed(all_keys):
            if key not in assessment_keys and key not in ['grade', 'grading', 'Grade', 'Grading']:
                attendance_column = key
                break
    
    # Headers for the attendance table: student info + attendance
    attendance_headers = student_info_keys.copy()
    if attendance_column and attendance_column not in attendance_headers:
        attendance_headers.append(attendance_column)

    # Create the attendance table
    table = doc.add_table(rows=1, cols=len(attendance_headers))

    # Set header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(attendance_headers):
        header_cells[i].text = header_text
        paragraph = header_cells[i].paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(header_text)
        run.bold = True
        run.font.size = Pt(9)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add student attendance rows
    for student in student_data:
        row_cells = table.add_row().cells
        
        # Add cell data for each header
        for i, header in enumerate(attendance_headers):
            value = student.get(header, '')
            row_cells[i].text = str(value)
        
        # Format cell text
        for cell in row_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)

    # Distribute width evenly across columns
    cell_width = Inches(7.0) / len(attendance_headers)
    widths = [cell_width] * len(attendance_headers)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = widths[idx]

    # Set table borders
    def set_cell_border(cell, border_type, border_size, border_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        element = tcPr.xpath(f"./w:{border_type}")
        if element:
            element[0].getparent().replace(element[0], border)
        else:
            tcPr.append(border)

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 3, '000000')
            set_cell_border(cell, 'bottom', 3, '000000')
            set_cell_border(cell, 'left', 3, '000000')
            set_cell_border(cell, 'right', 3, '000000')

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ensure table rows don't split across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)


################################################################  Attainment Table #########################

import matplotlib.pyplot as plt
import numpy as np
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_co_attainment_analysis(doc, data):
    """Generate CO attainment analysis tables with support for multi-page tables and charts"""
    if not all(key in data for key in ['coWeightages', 'studentData', 'coAttainmentCriteria', 'copoMappingData']):
        return  # Skip if required data is missing

    # Add section heading with page break
    doc.add_page_break()
    heading = doc.add_heading(level=1)
    run = heading.add_run('20. CO attainment analysis with the reflection on feedback on course outcomes')
    run.font.name = 'Carlito'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(28, 132, 196)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()

    # Calculate CO attainment
    result = calculate_co_attainment(data)
    cos = list(data.get('coWeightages', {}).keys())

    # 1. Generate CO Attainment Summary Table
    summary_heading = doc.add_heading('CO Attainment Summary', level=2)
    summary_heading.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    
    summary_table = doc.add_table(rows=6, cols=len(cos) + 1)
    doc.add_paragraph()
    
    # Header row
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Course Outcomes"
    
    for i, co in enumerate(cos):
        header_cells[i + 1].text = co
        format_cell(header_cells[i + 1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    format_cell(header_cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    # Data rows
    row_labels = ["Weights", "No. of students who scored at the highest attainment level (3)", 
                 "Percentage of students who scored at the highest attainment level (3)", "Attainment Level"]
    
    for row_idx, label in enumerate(row_labels):
        row = summary_table.rows[row_idx + 1]
        row.cells[0].text = label
        format_cell(row.cells[0], bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        for col_idx, co in enumerate(cos):
            cell = row.cells[col_idx + 1]
            
            if row_idx == 0:  # Weights
                cell.text = result["attainment_summary"]["weights"].get(co, "0.00%")
            elif row_idx == 1:  # Students scored 3
                cell.text = str(result["attainment_summary"]["students_scored3"].get(co, 0))
            elif row_idx == 2:  # Percentage scored 3
                cell.text = result["attainment_summary"]["percentage_scored3"].get(co, "0%")
            elif row_idx == 3:  # Attainment Level
                cell.text = str(result["attainment_summary"]["attainment_level"].get(co, 0))
            
            format_cell(cell, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Overall Course Attainment row
    overall_row = summary_table.rows[5]
    overall_row.cells[0].text = "Overall Course Attainment"
    format_cell(overall_row.cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    overall_row.cells[1].text = f"{result['attainment_summary']['overall_attainment']:.4f}"
    
    # Merge cells for overall attainment
    for i in range(2, len(cos) + 1):
        try:
            overall_row.cells[1].merge(overall_row.cells[i])
        except:
            pass
    
    format_cell(overall_row.cells[1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Add borders to summary table
    for row in summary_table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 4, '000000')
            set_cell_border(cell, 'bottom', 4, '000000')
            set_cell_border(cell, 'left', 4, '000000')
            set_cell_border(cell, 'right', 4, '000000')
    
    # Set narrower column widths similar to attendance report table
    # First column wider for labels, subsequent columns narrower for data
    col_widths = [Inches(2)] + [Inches(1)] * len(cos)
    set_table_column_widths(summary_table, col_widths)
    
    # Center align the table
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Make sure the table width is controlled
    summary_table.autofit = False
    
    # Reduce font size for better fit
    for row in summary_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Smaller font size for compact table
    
    # Make sure rows don't split across pages
    prevent_table_row_breaks(summary_table)
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add percentage scored ≥ 3 chart below the summary table
    add_percentage_scored3_chart(doc, result, cos)
    doc.add_paragraph()
    
    all_pos = set()
    for co, po_mappings in data.get('copoMappingData', {}).get('mappingData', {}).items():
        all_pos.update(po_mappings.keys())
    
    # Sort POs in a natural order (PO1, PO2, ..., PSO1, PSO2, ...)
    pos = sorted(list(all_pos), key=lambda x: (
        'PSO' in x, 
        int(''.join(filter(str.isdigit, x or '0'))) if any(c.isdigit() for c in x) else 0
    ))
    
    if pos:
        doc.add_page_break()
        program_heading = doc.add_heading('Program Attainment', level=2)
        program_heading.runs[0].font.size = Pt(14)
        doc.add_paragraph()
        
        program_table = doc.add_table(rows=2, cols=len(pos) + 1)
        
        # Header row
        header_cells = program_table.rows[0].cells
        header_cells[0].text = "Program Outcomes"
        
        for i, po in enumerate(pos):
            header_cells[i + 1].text = po
            format_cell(header_cells[i + 1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        format_cell(header_cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        # Data row
        data_row = program_table.rows[1]
        data_row.cells[0].text = "Program Attainment"
        format_cell(data_row.cells[0], bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        for col_idx, po in enumerate(pos):
            data_row.cells[col_idx + 1].text = result["program_attainment"].get(po, "0.00")
            format_cell(data_row.cells[col_idx + 1], bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Add borders to program table
        for row in program_table.rows:
            for cell in row.cells:
                set_cell_border(cell, 'top', 4, '000000')
                set_cell_border(cell, 'bottom', 4, '000000')
                set_cell_border(cell, 'left', 4, '000000')
                set_cell_border(cell, 'right', 4, '000000')
        
        program_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        prevent_table_row_breaks(program_table)
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add program attainment chart
        add_program_attainment_chart(doc, result, pos)
        doc.add_paragraph()
    
    # 3. Generate Student-wise CO Achievement Table on a new page
    if result["student_performance"]:
        doc.add_page_break()
        student_heading = doc.add_heading('Student-wise CO Achievement', level=2)
        student_heading.runs[0].font.size = Pt(14)
        doc.add_paragraph()
        
        # Create table with all students (will handle pagination automatically)
        rows_count = len(result["student_performance"]) + 2  # +2 for header and average row
        student_table = doc.add_table(rows=rows_count, cols=len(cos) + 1)
        
        # Header row
        header_cells = student_table.rows[0].cells
        header_cells[0].text = "NAME"
        
        for i, co in enumerate(cos):
            header_cells[i + 1].text = f"{co} Score"
            format_cell(header_cells[i + 1], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        format_cell(header_cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Student data rows
        for row_idx, student in enumerate(result["student_performance"]):
            row = student_table.rows[row_idx + 1]
            row.cells[0].text = student["rollNumber"]
            format_cell(row.cells[0], bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            for col_idx, co in enumerate(cos):
                cell = row.cells[col_idx + 1]
                cell.text = str(student["coScores"].get(co, ""))
                format_cell(cell, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Average row
        avg_row = student_table.rows[rows_count - 1]
        avg_row.cells[0].text = "Average"
        format_cell(avg_row.cells[0], bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        for col_idx, co in enumerate(cos):
            cell = avg_row.cells[col_idx + 1]
            cell.text = result["averages"].get(co, "0.00")
            format_cell(cell, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Set column widths for the student table
        set_table_column_widths(student_table, [Inches(2.0)] + [Inches(0.9)] * len(cos))
        
        # Add borders to student table
        for row in student_table.rows:
            for cell in row.cells:
                set_cell_border(cell, 'top', 4, '000000')
                set_cell_border(cell, 'bottom', 4, '000000')
                set_cell_border(cell, 'left', 4, '000000')
                set_cell_border(cell, 'right', 4, '000000')
        
        student_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc.add_paragraph()
        doc.add_paragraph()
        
        add_course_attainment_chart(doc, result, cos)
        doc.add_paragraph()

# Chart generation functions
def add_percentage_scored3_chart(doc, result, cos):
    """Create a horizontal bar chart showing Percentage of students who scored at the highest attainment level (3) for each CO"""
    plt.figure(figsize=(8, 4))
    
    # Extract data
    percentages = [float(result["attainment_summary"]["percentage_scored3"].get(co, "0%").replace("%", "")) for co in cos]
    
    # Create horizontal bar chart with proper color format
    y_pos = np.arange(len(cos))
    plt.barh(y_pos, percentages, color=(49/255, 85/255, 163/255, 0.6))  # Fixed color format
    
    # Customize chart
    plt.yticks(y_pos, cos)
    plt.xlabel('Percentage (%)')
    plt.title('Percentage of students who scored at the highest attainment level (3)')
    plt.xlim(0, 100)
    plt.grid(axis='x', linestyle='--', alpha=0.7)
    
    # Add value labels to the bars
    for i, v in enumerate(percentages):
        plt.text(v + 1, i, f"{v:.2f}%", va='center')
    
    # Save chart to memory and add to document
    image_stream = io.BytesIO()
    plt.tight_layout()
    plt.savefig(image_stream, format='png', dpi=100)
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(6))
    plt.close()
    
    # Center the image
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_program_attainment_chart(doc, result, pos):
    """Create a horizontal bar chart showing program attainment for each PO"""
    plt.figure(figsize=(8, 5))
    
    # Extract data
    attainments = [float(result["program_attainment"].get(po, "0.00")) for po in pos]
    
    # Create horizontal bar chart with proper color format
    y_pos = np.arange(len(pos))
    plt.barh(y_pos, attainments, color=(49/255, 85/255, 163/255, 0.6))  # Fixed color format
    
    # Customize chart
    plt.yticks(y_pos, pos)
    plt.xlabel('Attainment')
    plt.title('Program Attainment')
    plt.xlim(0, 3)
    plt.grid(axis='x', linestyle='--', alpha=0.7)
    
    # Add value labels to the bars
    for i, v in enumerate(attainments):
        plt.text(v + 0.1, i, f"{v:.2f}", va='center')
    
    # Save chart to memory and add to document
    image_stream = io.BytesIO()
    plt.tight_layout()
    plt.savefig(image_stream, format='png', dpi=100)
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(6))
    plt.close()
    
    # Center the image
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_course_attainment_chart(doc, result, cos):
    """Create a horizontal bar chart showing CO attainment with target line"""
    plt.figure(figsize=(8, 4))
    
    # Extract data
    attainments = [float(result["averages"].get(co, "0.00")) for co in cos]
    
    # Create horizontal bar chart with proper color format
    y_pos = np.arange(len(cos))
    plt.barh(y_pos, attainments, color=(49/255, 85/255, 163/255, 0.6), label='Attainment')
    
    # Add target line (value 3, which is the maximum score)
    plt.axvline(x=3, color=(237/255, 125/255, 49/255, 0.6), linestyle='-', linewidth=2, label='Required')
    
    # Customize chart
    plt.yticks(y_pos, cos)
    plt.xlabel('Attainment Score')
    plt.title('Course Outcome Attainment')
    plt.xlim(0, 3.5)
    plt.grid(axis='x', linestyle='--', alpha=0.7)
    plt.legend(loc='lower right')
    
    # Add value labels to the bars
    for i, v in enumerate(attainments):
        plt.text(v + 0.1, i, f"{v:.2f}", va='center')
    
    # Save chart to memory and add to document
    image_stream = io.BytesIO()
    plt.tight_layout()
    plt.savefig(image_stream, format='png', dpi=100)
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(6))
    plt.close()
    
    # Center the image
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Helper functions
def calculate_co_attainment(data):
    """Calculate CO attainment based on student data and weightages"""
    co_weightages = data.get('coWeightages', {})
    student_data = data.get('studentData', {})
    co_attainment_criteria = data.get('coAttainmentCriteria', {})
    copo_mapping_data = data.get('copoMappingData', {})
    target_attainment = data.get('targetAttainment', {})
    
    # Initialize results
    student_performance = []
    averages = {}
    program_attainment = {}
    attainment_summary = {
        "weights": {},
        "students_scored3": {},
        "percentage_scored3": {},
        "attainment_level": {},
        "overall_attainment": 0
    }
    
    # Get all CO keys
    cos = list(co_weightages.keys())
    
    # Calculate student performance and CO scores
    if student_data and 'maxMarks' in student_data and 'data' in student_data:
        # Get assessment components, excluding the total
        assessment_components = [(comp, max_mark) for comp, max_mark in student_data['maxMarks'].items() 
                                if 'Total' not in comp]
        
        # Process each student
        for student in student_data['data']:
            student_result = {
                "id": student.get("Unique Id."),
                "rollNumber": student.get("Student Name", ""),
                "coScores": {}
            }
            
            # Calculate CO scores for each student
            for co in cos:
                weighted_score = 0
                total_weight = 0
                
                for component, max_mark in assessment_components:
                    # Match component to co_weightages keys (case-insensitive)
                    component_key = component.lower()
                    
                    # Find matching key in co_weightages for this CO
                    co_component = None
                    for key in co_weightages.get(co, {}):
                        if component.lower().split('(')[0].strip() in key.lower():
                            co_component = key
                            break
                    
                    if co_component:
                        student_score = float(student.get(component, 0))
                        max_mark = float(max_mark if max_mark else 0)
                        co_weight = float(co_weightages[co].get(co_component, 0))
                        
                        weighted_score += (student_score * (co_weight / 100))
                        total_weight += (max_mark * (co_weight / 100))
                
                # Determine attainment level based on percentage
                partial = float(co_attainment_criteria.get(co, {}).get('partial', 0))
                full = float(co_attainment_criteria.get(co, {}).get('full', 0))
                
                percentage = (weighted_score / total_weight) * 100 if total_weight > 0 else 0
                
                # Assign score based on student's percentage
                if percentage >= full:
                    student_result['coScores'][co] = 3
                elif percentage >= partial:
                    student_result['coScores'][co] = 2
                else:
                    student_result['coScores'][co] = 1
            
            student_performance.append(student_result)
    
    # Calculate CO weights
    attainment_summary["weights"] = calculate_co_weights(co_weightages, student_data)
    
    # Calculate averages and attainment summary
    for co in cos:
        scores = [student['coScores'].get(co, 0) for student in student_performance]
        avg = sum(scores) / len(scores) if scores else 0
        averages[co] = f"{avg:.2f}"
        
        # Count students who scored 3
        scored3_count = sum(1 for score in scores if score >= 3)
        attainment_summary["students_scored3"][co] = scored3_count
        
        # Calculate percentage of students who scored 3
        percentage_scored3 = (scored3_count / len(scores)) * 100 if scores else 0
        attainment_summary["percentage_scored3"][co] = f"{percentage_scored3:.2f}%"
        
        # Calculate attainment level based on percentage of students who scored 3
        # This is critical for determining the overall attainment
        if target_attainment and co in target_attainment:
            full = float(target_attainment[co].get('full', 0))
            partial = float(target_attainment[co].get('partial', 0))
            
            if percentage_scored3 >= full:
                attainment_summary["attainment_level"][co] = 3
            elif percentage_scored3 >= partial:
                attainment_summary["attainment_level"][co] = 2
            else:
                attainment_summary["attainment_level"][co] = 1
        else:
            # Fallback to criteria from coAttainmentCriteria if targetAttainment is not available
            if co_attainment_criteria and co in co_attainment_criteria:
                full = float(co_attainment_criteria[co].get('full', 0))
                partial = float(co_attainment_criteria[co].get('partial', 0))
                
                if percentage_scored3 >= full:
                    attainment_summary["attainment_level"][co] = 3
                elif percentage_scored3 >= partial:
                    attainment_summary["attainment_level"][co] = 2
                else:
                    attainment_summary["attainment_level"][co] = 1
            else:
                attainment_summary["attainment_level"][co] = 1
    
    # Calculate overall attainment as average of individual CO attainments
    attainment_values = list(attainment_summary["attainment_level"].values())
    attainment_summary["overall_attainment"] = sum(attainment_values) / len(attainment_values) if attainment_values else 0
    
    # Calculate program attainment
    program_attainment = calculate_program_attainment(copo_mapping_data, averages)
    
    return {
        "student_performance": student_performance,
        "averages": averages,
        "program_attainment": program_attainment,
        "attainment_summary": attainment_summary
    }

def calculate_co_weights(co_weightages, student_data):
    """Calculate CO weights based on assessment components"""
    cos = list(co_weightages.keys())
    weights = {}
    
    # Guard against studentData or maxMarks being missing
    if not student_data or 'maxMarks' not in student_data:
        return {co: "0.00%" for co in cos}
    
    assessment_components = [(comp, mark) for comp, mark in student_data['maxMarks'].items() 
                            if 'Total' not in comp]
    total_max_marks = sum(float(mark or 0) for _, mark in assessment_components)
    
    for co in cos:
        co_weighted_sum = 0
        
        for component, max_mark in assessment_components:
            # Match component to co_weightages keys (case-insensitive)
            co_component = None
            for key in co_weightages.get(co, {}):
                if component.lower().split('(')[0].strip() in key.lower():
                    co_component = key
                    break
            
            if co_component:
                co_weight = float(co_weightages[co].get(co_component, 0))
                co_weighted_sum += (float(max_mark or 0) * (co_weight / 100))
        
        # Calculate percentage weight
        weight_percentage = (co_weighted_sum / total_max_marks * 100) if total_max_marks > 0 else 0
        weights[co] = f"{weight_percentage:.2f}%"
    
    return weights

def calculate_program_attainment(copo_mapping_data, averages):
    """Calculate program attainment based on CO-PO mapping and CO averages"""
    program_attainments = {}
    weight_sums = {}
    
    # Get all possible POs from the mapping data
    all_pos = set()
    for co, po_mappings in copo_mapping_data.get('mappingData', {}).items():
        all_pos.update(po_mappings.keys())
    
    # Initialize program attainments and weight sums for all POs
    for po in all_pos:
        program_attainments[po] = 0
        weight_sums[po] = 0
    
    # Calculate weighted sums for each PO
    for co, po_mappings in copo_mapping_data.get('mappingData', {}).items():
        for po, mapping_value in po_mappings.items():
            # Convert mapping value to float, treat empty string as 0
            mapping_value = float(mapping_value) if mapping_value and mapping_value.strip() else 0
            co_average = float(averages.get(co, 0))
            
            # Sum up (mapping_value * CO_average) for each PO
            program_attainments[po] += mapping_value * co_average
            # Keep track of total mapping values for each PO
            weight_sums[po] += mapping_value
    
    # Calculate final weighted averages
    for po in program_attainments:
        if weight_sums[po] > 0:
            program_attainments[po] = f"{program_attainments[po] / weight_sums[po]:.2f}"
        else:
            program_attainments[po] = "0.00"
    
    return program_attainments

def format_cell(cell, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Format a table cell"""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    
    for run in paragraph.runs:
        run.font.bold = bold
        run.font.size = Pt(10)
    
    # If no runs, create one
    if not paragraph.runs:
        run = paragraph.add_run(cell.text)
        run.font.bold = bold
        run.font.size = Pt(10)
        cell.text = ""  # Clear text since we've added it to the run

def set_cell_border(cell, border_type, border_size, border_color):
    """Set cell border properties"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    border = OxmlElement(f'w:{border_type}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(border_size))
    border.set(qn('w:color'), border_color)
    element = tcPr.xpath(f"./w:{border_type}")
    if element:
        element[0].getparent().replace(element[0], border)
    else:
        tcPr.append(border)

def set_table_column_widths(table, widths):
    """Set column widths for a table"""
    for i, width in enumerate(widths):
        for row in table.rows:
            if i < len(row.cells):
                row.cells[i].width = width

def prevent_table_row_breaks(table):
    """Ensure table rows don't split across pages"""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

create_co_attainment_analysis(doc, data)

########################################################################################################################

# Feedback Data Section
def create_feedback_section(doc, data):
    if data.get('feedbackData'):
        # Add page break and section heading
        doc.add_page_break()
        feedback_heading = doc.add_heading(level=1)
        feedback_run = feedback_heading.add_run('21. Feedback (class committee or otherwise) and corrective actions (if any)')
        feedback_run.font.name = 'Carlito'
        feedback_run.font.size = Pt(16)
        feedback_run.font.color.rgb = RGBColor(28, 132, 196)
        feedback_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add quantitative feedback
        if data['feedbackData'].get('quantitativeFeedback'):
            quant_heading = doc.add_paragraph()
            quant_heading_run = quant_heading.add_run('Quantitative Feedback:')
            quant_heading_run.bold = True
            quant_heading_run.font.size = Pt(12)
            quant_heading.paragraph_format.left_indent = Inches(0.7)
            
            quant_para = doc.add_paragraph()
            quant_para.paragraph_format.left_indent = Inches(0.7)
            quant_para.paragraph_format.right_indent = Inches(0.5)
            quant_run = quant_para.add_run(f"Average Rating: {data['feedbackData']['quantitativeFeedback']}/5")
            quant_run.font.size = Pt(12)
        
        # Add qualitative feedback
        if data['feedbackData'].get('qualitativeFeedback'):
            doc.add_paragraph()  # Add space

            # Heading
            qual_heading = doc.add_paragraph()
            qual_heading_run = qual_heading.add_run('Qualitative Feedback:')
            qual_heading_run.bold = True
            qual_heading_run.font.size = Pt(12)
            qual_heading.paragraph_format.left_indent = Inches(0.7)

            # Feedback content
            feedback_text = data['feedbackData']['qualitativeFeedback'].encode('utf-8', 'ignore').decode('utf-8')
            for line in feedback_text.split('\n'):
                if line.strip():  # skip empty lines
                    para = doc.add_paragraph(line.strip())
                    para.paragraph_format.left_indent = Inches(0.7)
                    para.paragraph_format.right_indent = Inches(0.5)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = para.runs[0]
                    run.font.size = Pt(12)

############################################################################################################

# Faculty Course Review Section
def create_faculty_review_section(doc, data):
    if data.get('facultyCourseReview'):
        # Add page break and section heading
        doc.add_page_break()
        review_heading = doc.add_heading(level=1)
        review_run = review_heading.add_run('22. Faculty Course Review')
        review_run.font.name = 'Carlito'
        review_run.font.size = Pt(16)
        review_run.font.color.rgb = RGBColor(28, 132, 196)
        review_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Process review content
        review_text = data['facultyCourseReview'].encode('utf-8', 'ignore').decode('utf-8')
        for line in review_text.split('\n'):
            if line.strip():  # Skip empty lines
                para = doc.add_paragraph(line.strip())
                para.paragraph_format.left_indent = Inches(0.7)
                para.paragraph_format.right_indent = Inches(0.5)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = para.runs[0]
                run.font.size = Pt(12)

# Call these functions after creating actions doc but before saving the document
create_feedback_section(doc, data)
create_faculty_review_section(doc, data)

#######################################################################################################################
# Saving the Document and Converting to PDF
unique_suffix = uuid.uuid4().hex
output_doc = './download/' + data['filename'][:-4] + "_" + unique_suffix + '.docx'
doc.save(output_doc)
print("Document updated and saved:", output_doc)

# Use a file lock to serialize the conversion process
lock_path = "./download/conversion.lock"
lock = FileLock(lock_path, timeout=60)  # waits up to 60 seconds for the lock
try:
    with lock:
        print("Acquired lock for conversion.")
        try:
            convert(output_doc)
        except Exception as e:
            print("Conversion error:", e)
            sys.exit(1)
        print("Conversion completed for:", output_doc)
except Timeout:
    print("Could not acquire lock for conversion. Another process may be converting.")
    sys.exit(1)

# The PDF is generated with the same unique suffix
pdf_path = output_doc.replace('.docx', '.pdf')

# Remove the DOCX file after conversion
try:
    os.remove(output_doc)
    print("Removed temporary DOCX:", output_doc)
except Exception as e:
    print("Error removing DOCX file:", e)

# Prepare PDF list for merging
# If an assignmentPDF is provided, include it; otherwise, just use our generated PDF.

from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO

def add_heading_to_pdf(input_pdf_path, heading_text):
    # Step 1: Create a PDF with the heading text
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.setFont("Helvetica-Bold", 14)
    can.drawString(72, 750, heading_text)  # 1 inch margin from left and near the top
    can.save()
    packet.seek(0)

    # Step 2: Read original PDF and overlay the heading
    heading_pdf = PdfReader(packet)
    original_pdf = PdfReader(input_pdf_path)
    writer = PdfWriter()

    # Merge heading on top of the first page
    first_page = original_pdf.pages[0]
    first_page.merge_page(heading_pdf.pages[0])
    writer.add_page(first_page)

    # Add the rest of the pages unchanged
    for page in original_pdf.pages[1:]:
        writer.add_page(page)

    # Step 3: Save modified version (overwrite original or use temp path)
    modified_pdf_path = input_pdf_path.replace(".pdf", "_with_heading.pdf")
    with open(modified_pdf_path, "wb") as f_out:
        writer.write(f_out)

    return modified_pdf_path

# if data.get('assignmentPDF'):
#     assignment_pdf_path = "./data/assignments/" + data['assignmentPDF']
#     heading_text = "Assignments / Quiz / Internal Components"
#     updated_assignment_pdf = add_heading_to_pdf(assignment_pdf_path, heading_text)
#     pdf_list = [pdf_path, updated_assignment_pdf]
# else:
#     pdf_list = [pdf_path]

if data.get('assignmentPDF'):
    pdf_list = [pdf_path, "./data/assignments/" + data['assignmentPDF']]
else:
    pdf_list = [pdf_path]

merger = PdfMerger()
for pdf in pdf_list:
    print("Appending PDF:", pdf)
    merger.append(pdf)

# Merge PDFs into the final compiled PDF
final_pdf_path = "./download/" + data['filename']
merger.write(final_pdf_path)
merger.close()
print("Final PDF generated:", final_pdf_path)

# Cleanup: Remove any and every other file (temporary PDFs) that are not the final compiled PDF.
for pdf in pdf_list:
    # Only remove temporary files that are in the './download' folder.
    if pdf != final_pdf_path and pdf.startswith("./download/") and os.path.exists(pdf):
        try:
            os.remove(pdf)
            print("Removed temporary PDF:", pdf)
        except Exception as e:
            print("Error removing temporary PDF file:", e)

# Optionally: Remove the lock file if it exists (this is typically cleaned up automatically)
if os.path.exists(lock_path):
    try:
        os.remove(lock_path)
        print("Removed lock file:", lock_path)
    except Exception as e:
        print("Error removing lock file:", e)